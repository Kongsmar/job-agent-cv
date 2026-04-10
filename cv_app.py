import streamlit as st
from openai import OpenAI
import io
import requests
from bs4 import BeautifulSoup
from docx import Document
from PyPDF2 import PdfReader
import json
import re

# --- KONFIGURATION & DESIGN ---
st.set_page_config(page_title="Strategisk CV-Builder Pro", page_icon="🎓", layout="wide")

st.markdown("""
<style>
    .stApp { background-color: #1a1c24; color: #e0e0e0; }
    h1 { text-align: center; color: #ffffff !important; border-bottom: 2px solid #4a90e2; padding-bottom: 15px; font-weight: 800; }
    
    /* CV Sektioner */
    .cv-block {
        background-color: #2d303d;
        padding: 25px;
        border-radius: 12px;
        border-left: 5px solid #4a90e2;
        margin-bottom: 25px;
        box-shadow: 0 6px 10px rgba(0, 0, 0, 0.4);
    }
    
    .cv-section-title {
        font-size: 1.4em;
        font-weight: bold;
        color: #4a90e2;
        margin-bottom: 15px;
        border-bottom: 1px solid #4a4d5e;
        padding-bottom: 5px;
        text-transform: uppercase;
    }
    
    .cv-text {
        font-family: 'Segoe UI', sans-serif;
        line-height: 1.6;
        white-space: pre-wrap;
        color: #f0f0f0;
    }

    .entry-headline {
        background-color: #262936;
        padding: 8px 12px;
        border-radius: 6px;
        border: 1px solid #3a3d4d;
        color: #ffffff;
        font-weight: bold;
        margin-bottom: 8px;
    }

    /* Analyse visning */
    .analyse-block {
        background-color: #262936;
        padding: 25px;
        border-radius: 12px;
        border: 2px solid #4a90e2;
        margin-bottom: 30px;
    }

    .score-container {
        background-color: #0e1117;
        padding: 15px;
        border-radius: 10px;
        text-align: center;
        border: 1px solid #4a90e2;
    }
    
    .score-number {
        font-size: 2.8em;
        font-weight: 800;
        color: #4a90e2;
        display: block;
    }
</style>
""", unsafe_allow_html=True)

# --- HJÆLPEFUNKTIONER ---
def get_text_from_url(url):
    try:
        headers = {'User-Agent': 'Mozilla/5.0'}
        response = requests.get(url, headers=headers, timeout=10)
        response.encoding = 'utf-8'
        soup = BeautifulSoup(response.text, 'html.parser')
        for element in soup(["script", "style", "nav", "footer"]): element.extract()
        return "\n".join([tag.get_text().strip() for tag in soup.find_all(['h1', 'h2', 'p', 'li']) if tag.get_text()])
    except: return "Kunne ikke hente tekst."

def extract_pdf(file):
    try:
        reader = PdfReader(file)
        return "".join([p.extract_text() for p in reader.pages])
    except: return ""

def format_text_for_word(text):
    if not text or len(str(text).strip()) < 2: return ""
    # Gør [OVERSKRIFT] pænere i Word
    formatted = re.sub(r'\[(.*?)\]', r'\n\n\1\n', str(text))
    return formatted.strip()

def fill_cv_docx(template, data_dict):
    try:
        template.seek(0)
        doc = Document(template)
        clean_data = {k: format_text_for_word(v) for k, v in data_dict.items()}
        
        for p in doc.paragraphs:
            for key, value in clean_data.items():
                if key in p.text: p.text = p.text.replace(key, value)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for key, value in clean_data.items():
                            if key in p.text: p.text = p.text.replace(key, value)
        
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf
    except: return None

def style_cv_entries(raw_text):
    if not raw_text or len(str(raw_text).strip()) < 5: return None
    pattern = r'(\[.*?\])'
    segments = re.split(pattern, str(raw_text))
    formatted_html = ""
    current_headline = None
    for segment in segments:
        segment = segment.strip()
        if not segment: continue
        if segment.startswith('[') and segment.endswith(']'):
            current_headline = segment.replace('[', '').replace(']', '')
        else:
            if current_headline:
                formatted_html += f"""
                <div style='margin-bottom: 20px;'>
                    <div class='entry-headline'>{current_headline}</div>
                    <div class='cv-text'>{segment}</div>
                </div>"""
                current_headline = None
            else:
                formatted_html += f"<div style='margin-bottom:15px;' class='cv-text'>{segment}</div>"
    return formatted_html

# --- APP FLOW ---
st.markdown("<h1>🎓 Strategisk & Akademisk CV-Builder</h1>", unsafe_allow_html=True)

if 'cv_step' not in st.session_state: st.session_state.cv_step = 1

if st.session_state.cv_step == 1:
    col1, col2 = st.columns(2, gap="large")
    with col1:
        st.subheader("1. Dit Grundlag")
        master_cv = st.file_uploader("Upload dit Master-CV (PDF)", type="pdf")
        cv_template = st.file_uploader("Upload din Word-skabelon (DOCX)", type="docx")
        navn = st.text_input("Dit fulde navn:")
    with col2:
        st.subheader("2. Jobmålet")
        job_url = st.text_input("Link til jobopslag (valgfrit):")
        if st.button("Hent jobtekst fra link 🌐") and job_url:
            st.session_state.temp_job_text = get_text_from_url(job_url)
        job_text = st.text_area("Indsæt jobbeskrivelsen her:", value=st.session_state.get('temp_job_text', ""), height=250)

    if st.button("Generér strategisk CV ✨", type="primary", use_container_width=True):
        if master_cv and job_text:
            st.session_state.master_cv_text = extract_pdf(master_cv)
            st.session_state.job_content = job_text
            st.session_state.cv_template = cv_template
            st.session_state.user_name = navn
            st.session_state.cv_step = 2
            st.rerun()
        else:
            st.warning("Husk at uploade dit Master-CV og indsætte en jobtekst.")

elif st.session_state.cv_step == 2:
    with st.spinner("AI optimerer dit indhold efter strategiske og akademiske principper..."):
        try:
            client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
            prompt = f"""
            Du er en elite-rekrutteringskonsulent. Omskriv Master-CV'et baseret på disse regler:
            
            1. RELEVANS: Fjern alt irrelevant. Spejl virksomhedens tone og ordvalg fra jobopslaget.
            2. RESULTATER: Brug formlen: "Jeg har opnået [reelt resultat] ved at anvende [kompetence]".
            3. SANDFÆRDIGHED: Du må IKKE opfinde tal, resultater eller procenter, der ikke findes i Master-CV'et.
            4. AKADEMISK FOKUS: Fremhæv relevante fag, metoder og specialiseringer. Gør viden anvendelig.
            5. PROFILTEKST: 5-8 linjer øverst (elevatortale). Skal være 100% faglig og målrettet.
            6. STRUKTUR: Omvendt kronologisk. Hver post starter med [TITEL | VIRKSOMHED | PERIODE].
            7. TOMME SEKTIONER: Hvis en sektion (f.eks. 'kurser' eller 'fritid') ikke findes i Master-CV'et, returner "".

            SVAR KUN I JSON FORMAT:
            {{
              "analyse": {{ "score": int, "vurdering": str }},
              "kontakt": str, "profil": str, "erfaring": str, "uddannelse": str, 
              "kompetencer": str, "kurser": str, "sprog": str, "fritid": str
            }}
            DATA: JOB: {st.session_state.job_content} | Master-CV: {st.session_state.master_cv_text}
            """
            
            resp = client.chat.completions.create(model="gpt-4o", messages=[{"role": "user", "content": prompt}], response_format={"type": "json_object"})
            res = json.loads(resp.choices[0].message.content)
            ana = res.get('analyse', {})

            # --- ANALYSE VISNING ---
            st.markdown("<div class='analyse-block'>", unsafe_allow_html=True)
            col_score, col_details = st.columns([1, 2])
            with col_score:
                st.markdown(f"<div class='score-container'><span class='score-number'>{ana.get('score')}%</span>Match</div>", unsafe_allow_html=True)
            with col_details:
                st.markdown(f"### Strategisk Vurdering")
                st.write(f"{ana.get('vurdering')}")
                st.progress(ana.get('score', 0) / 100)
            st.markdown("</div>", unsafe_allow_html=True)

            # --- FORHÅNDSVISNING ---
            st.markdown(f"<div class='cv-block' style='text-align:center;'><h1>{st.session_state.user_name}</h1>{res.get('kontakt', '')}</div>", unsafe_allow_html=True)
            
            col_l, col_r = st.columns([2, 1], gap="medium")
            with col_l:
                if res.get('profil'):
                    st.markdown(f"<div class='cv-block'><div class='cv-section-title'>Profil</div><div class='cv-text'>{res.get('profil')}</div></div>", unsafe_allow_html=True)
                
                erfaring_html = style_cv_entries(res.get('erfaring'))
                if erfaring_html:
                    st.markdown(f"<div class='cv-block'><div class='cv-section-title'>Erhvervserfaring & Resultater</div>{erfaring_html}</div>", unsafe_allow_html=True)
            
            with col_r:
                if res.get('kompetencer'):
                    st.markdown(f"<div class='cv-block'><div class='cv-section-title'>Kompetencer</div><div class='cv-text'>{res.get('kompetencer')}</div></div>", unsafe_allow_html=True)
                
                uddannelse_html = style_cv_entries(res.get('uddannelse'))
                if uddannelse_html:
                    st.markdown(f"<div class='cv-block'><div class='cv-section-title'>Uddannelse</div>{uddannelse_html}</div>", unsafe_allow_html=True)
                
                kurser_html = style_cv_entries(res.get('kurser'))
                if kurser_html:
                    st.markdown(f"<div class='cv-block'><div class='cv-section-title'>Kurser</div>{kurser_html}</div>", unsafe_allow_html=True)
                
                if res.get('sprog'):
                    st.markdown(f"<div class='cv-block'><div class='cv-section-title'>Sprog</div><div class='cv-text'>{res.get('sprog')}</div></div>", unsafe_allow_html=True)
                
                if res.get('fritid'):
                    st.markdown(f"<div class='cv-block'><div class='cv-section-title'>Personligt / Fritid</div><div class='cv-text'>{res.get('fritid')}</div></div>", unsafe_allow_html=True)

            # --- DOWNLOAD ---
            if st.session_state.cv_template:
                replacements = {
                    "{{NAVN}}": st.session_state.user_name,
                    "{{CV_KONTAKT}}": res.get('kontakt', ''),
                    "{{CV_PROFIL}}": res.get('profil', ''),
                    "{{CV_ERFARING}}": res.get('erfaring', ''),
                    "{{CV_UDDANNELSE}}": res.get('uddannelse', ''),
                    "{{CV_KURSER}}": res.get('kurser', ''),
                    "{{CV_SPROG}}": res.get('sprog', ''),
                    "{{CV_KOMPETENCER}}": res.get('kompetencer', ''),
                    "{{CV_FRITID}}": res.get('fritid', '')
                }
                final_doc = fill_cv_docx(st.session_state.cv_template, replacements)
                st.download_button("Download målrettet CV (.docx) 📄", final_doc, f"CV_{st.session_state.user_name}.docx", type="primary", use_container_width=True)

        except Exception as e:
            st.error(f"Der skete en fejl: {e}")

    if st.button("Start forfra 🔄"):
        for key in list(st.session_state.keys()): del st.session_state[key]
        st.session_state.cv_step = 1
        st.rerun()
